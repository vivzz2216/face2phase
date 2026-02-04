"""
PDF and Word Report Generator with Charts and Analytics
"""
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
import json

# PDF Generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

# Word Document Generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Charts
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import logging as mpl_logging
mpl_logging.getLogger('matplotlib.font_manager').setLevel(mpl_logging.WARNING)  # Suppress font manager debug logs
import seaborn as sns
import numpy as np

from ..core.settings import REPORTS_DIR, VALIDATION_ENABLED

# Import validation utilities for data integrity
try:
    from ..utils.validators import metric_validator, validate_analysis_results
    VALIDATORS_AVAILABLE = True
except ImportError:
    VALIDATORS_AVAILABLE = False

logger = logging.getLogger(__name__)

class PDFWordReportGenerator:
    """Generate professional PDF and Word reports with analytics"""
    
    # Standard rounding precision for all percentages/scores
    DECIMAL_PRECISION = 2
    
    def __init__(self):
        self.reports_dir = REPORTS_DIR
        self.charts_dir = self.reports_dir / "charts"
        self.charts_dir.mkdir(exist_ok=True)
    
    def _validate_and_prepare_data(self, analysis_data: Dict) -> Dict:
        """
        Validate analysis data before PDF generation to ensure accuracy
        
        - Validates all scores are within valid ranges (0-100)
        - Ensures mathematical consistency
        - Rounds all values to standard precision
        - Returns corrected data if issues found
        """
        validated_data = dict(analysis_data)
        validation_warnings = []
        
        if VALIDATORS_AVAILABLE and VALIDATION_ENABLED:
            try:
                is_valid, corrected, errors = validate_analysis_results(analysis_data)
                if not is_valid:
                    for error in errors:
                        logger.warning(f"PDF data validation issue: {error}")
                        validation_warnings.append(error)
                    validated_data = corrected
            except Exception as e:
                logger.warning(f"Could not validate analysis data: {e}")
        
        # Ensure all score values are properly rounded and clamped
        score_keys = ['overall_score', 'voice_confidence', 'facial_confidence', 'vocabulary_score']
        for key in score_keys:
            if key in validated_data:
                try:
                    value = float(validated_data[key])
                    # Clamp to valid range [0, 100]
                    value = max(0.0, min(100.0, value))
                    # Round to standard precision
                    validated_data[key] = round(value, self.DECIMAL_PRECISION)
                except (TypeError, ValueError):
                    logger.warning(f"Invalid score value for {key}: {validated_data.get(key)}")
                    validated_data[key] = 0.0
        
        # Validate filler analysis percentages
        if 'filler_analysis' in validated_data:
            filler = validated_data.get('filler_analysis', {})
            if 'filler_percentage' in filler:
                try:
                    pct = float(filler['filler_percentage'])
                    pct = max(0.0, min(100.0, pct))
                    filler['filler_percentage'] = round(pct, self.DECIMAL_PRECISION)
                except (TypeError, ValueError):
                    filler['filler_percentage'] = 0.0
        
        # Store validation warnings for potential display
        validated_data['_validation_warnings'] = validation_warnings
        
        return validated_data
    
    def generate_pdf_report(self, session_id: str, analysis_data: Dict) -> Optional[Path]:
        """
        Generate comprehensive PDF report with charts
        
        Args:
            session_id: Session identifier
            analysis_data: Complete analysis results
            
        Returns:
            Path to generated PDF file or None if generation fails
        """
        try:
            # VALIDATION: Ensure data integrity before PDF generation
            analysis_data = self._validate_and_prepare_data(analysis_data)
            
            # Log any validation warnings
            warnings = analysis_data.get('_validation_warnings', [])
            if warnings:
                logger.info(f"PDF generation proceeding with {len(warnings)} validation corrections")
            
            # Ensure reports directory exists
            self.reports_dir.mkdir(exist_ok=True)
            pdf_path = self.reports_dir / f"{session_id}_report.pdf"
            
            # Create PDF document with enhanced margins
            doc = SimpleDocTemplate(
                str(pdf_path),
                pagesize=letter,
                rightMargin=60,
                leftMargin=60,
                topMargin=50,
                bottomMargin=40
            )
            
            # Container for PDF elements
            story = []
            styles = getSampleStyleSheet()
            
            # === MODERN PROFESSIONAL STYLES ===
            # Brand Colors
            primary_color = colors.HexColor('#0f172a')  # Slate 900
            accent_color = colors.HexColor('#3b82f6')   # Blue 500
            success_color = colors.HexColor('#10b981')  # Emerald 500
            bg_light = colors.HexColor('#f8fafc')       # Slate 50
            
            # Typography
            styles = getSampleStyleSheet()
            
            # Title: Big, Bold, Dark
            title_style = ParagraphStyle(
                'ModernTitle',
                parent=styles['Heading1'],
                fontSize=24,
                fontName='Helvetica-Bold',
                textColor=primary_color,
                spaceAfter=4,
                alignment=0 # Left aligned
            )
            
            # Subtitle: Uppercase tracking
            subtitle_style = ParagraphStyle(
                'ModernSubtitle',
                parent=styles['Normal'],
                fontSize=10,
                fontName='Helvetica',
                textColor=colors.HexColor('#64748b'), # Slate 500
                spaceAfter=20,
                alignment=0
            )
            
            # Section Headers: Clean, partial underline look via border (simulated)
            heading_style = ParagraphStyle(
                'ModernHeading',
                parent=styles['Heading2'],
                fontSize=14,
                fontName='Helvetica-Bold',
                textColor=primary_color,
                spaceBefore=24,
                spaceAfter=12,
            )
            
            # Normal Text: Readable size, good leading
            normal_style = ParagraphStyle(
                'ModernNormal',
                parent=styles['Normal'],
                fontSize=10,
                fontName='Helvetica',
                textColor=colors.HexColor('#334155'), # Slate 700
                leading=14
            )
            
            # Label Text (for small labels)
            label_style = ParagraphStyle(
                'ModernLabel',
                parent=styles['Normal'],
                fontSize=8,
                fontName='Helvetica-Bold',
                textColor=colors.HexColor('#64748b'),
                alignment=1 # Center
            )

            # === HEADER SECTION ===
            # Two columns: Brand Left | Meta Right
            timestamp = analysis_data.get('timestamp', datetime.now().isoformat())
            try:
                dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                date_str = dt.strftime('%B %d, %Y')
                time_str = dt.strftime('%I:%M %p')
            except:
                date_str = timestamp
                time_str = ""

            header_data = [
                [Paragraph("FACE2PHASE", title_style), Paragraph(f"SESSION REPORT<br/>{date_str} • {time_str}", 
                 ParagraphStyle('MetaRight', parent=subtitle_style, alignment=2))]
            ]
            header_table = Table(header_data, colWidths=[4*inch, 3*inch])
            header_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(header_table)
            
            story.append(Spacer(1, 4))
            # Separator Line
            story.append(Table([['']], colWidths=[7*inch], style=TableStyle([
                ('LINEBELOW', (0,0), (-1,-1), 1, colors.HexColor('#e2e8f0')),
                 ('LEFTPADDING', (0,0), (-1,-1), 0),
            ])))
            story.append(Spacer(1, 25))

            # === KPI METRICS GRID (SCORE CARDS) ===
            # We will create a grid of 4 boxes for the top level metrics
            
            overall_score = analysis_data.get('overall_score', 0)
            voice_score = analysis_data.get('voice_confidence', 0)
            vocab_score = analysis_data.get('vocabulary_score', 0)
            facial_score = analysis_data.get('facial_confidence', 0)
            is_audio_only = analysis_data.get('is_audio_only', False)

            def score_cell(title, value, suffix="/100", color_override=None):
                val_float = float(value)
                color = color_override or (success_color if val_float >= 70 else colors.HexColor('#f59e0b') if val_float >= 50 else colors.HexColor('#ef4444'))
                
                # Big Number Style
                num_style = ParagraphStyle(
                    'ScoreNum', parent=normal_style, fontSize=24, fontName='Helvetica-Bold', 
                    textColor=color, alignment=1, spaceAfter=2
                )
                
                return [
                    Paragraph(title, label_style),
                    Paragraph(f"{val_float:.0f}{suffix}", num_style)
                ]

            # Row of cards
            card_data = [
                score_cell("OVERALL SCORE", overall_score),
                score_cell("VOICE CONFIDENCE", voice_score),
                score_cell("VOCABULARY", vocab_score),
            ]
            
            col_widths = [1.6*inch, 1.6*inch, 1.6*inch]
            
            if not is_audio_only:
                card_data.append(score_cell("FACIAL CONFIDENCE", facial_score))
                col_widths.append(1.6*inch)
            else:
                # Fill space if 3 cols
                col_widths = [2.2*inch, 2.2*inch, 2.2*inch]

            # Container Table for cards
            # We want gaps between cards, so we use a trick (spacer columns) or just padding
            # Simpler with styling:
            
            cards_table = Table([card_data], colWidths=col_widths)
            cards_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                # Add background and borders to make them look like cards
                ('BACKGROUND', (0,0), (-1,-1), bg_light),
                ('ROUNDEDCORNERS', [4, 4, 4, 4]), # Attempt rounded corners if supported or just separate cells
                ('TOPPADDING', (0,0), (-1,-1), 12),
                ('BOTTOMPADDING', (0,0), (-1,-1), 12),
                ('GRID', (0,0), (-1,-1), 1, colors.white), # White grid to separate "cards" visually on the gray bg
            ]))
            
            story.append(cards_table)
            story.append(Spacer(1, 25))
            
            # Strict Evaluation Breakdown (if available)
            if 'strict_evaluation' in analysis_data:
                strict_eval = analysis_data.get('strict_evaluation') or {}
                scores = strict_eval.get('scores', {})
                
                # Check if we have valid non-zero scores
                valid_scores = any(scores.get(k, 0) > 0 for k in ['clarity_pronunciation_25', 'fluency_pace_25', 'coherence_grammar_25', 'content_accuracy_25'])
                
                breakdown_data = []
                if valid_scores:
                    story.append(Paragraph("Detailed Score Breakdown", heading_style))
                    
                    breakdown_data = [['Dimension', 'Points', 'Score', 'Max']]
                    
                    # Add each score only if it exists
                    if 'clarity_pronunciation_25' in scores:
                        breakdown_data.append(['Clarity & Pronunciation', '25', f"{scores.get('clarity_pronunciation_25', 0):.1f}", '25.0'])
                    if 'fluency_pace_20' in scores or 'fluency_pace_25' in scores:
                        # Handle both old and new keys for backward compatibility
                        val = scores.get('fluency_pace_25') or scores.get('fluency_pace_20', 0)
                        breakdown_data.append(['Fluency & Pace', '25', f"{val:.1f}", '25.0'])
                    if 'coherence_grammar_25' in scores:
                        breakdown_data.append(['Coherence & Grammar', '25', f"{scores.get('coherence_grammar_25', 0):.1f}", '25.0'])
                    if 'content_accuracy_20' in scores or 'content_accuracy_25' in scores:
                        val = scores.get('content_accuracy_25') or scores.get('content_accuracy_20', 0)
                        breakdown_data.append(['Content Accuracy', '25', f"{val:.1f}", '25.0'])
                    if 'delivery_engagement_10' in scores:
                        breakdown_data.append(['Delivery & Engagement', '10', f"{scores.get('delivery_engagement_10', 0):.1f}", '10.0'])
                    if 'final_100' in scores:
                        breakdown_data.append(['Total Score', '100', f"{scores.get('final_100', 0):.1f}", '100.0'])

                if len(breakdown_data) > 1:
                    breakdown_table = Table(breakdown_data, colWidths=[2.2*inch, 1.2*inch, 1.2*inch, 1.2*inch])
                    
                    # Find the last row index for total score highlighting
                    last_row_idx = len(breakdown_data) - 1
                    
                    breakdown_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), primary_color),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 11),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                        ('TOPPADDING', (0, 0), (-1, 0), 10),
                        ('BACKGROUND', (0, 1), (-1, -1), bg_light),
                        ('LINEBELOW', (0, 0), (-1, -2), 1, colors.HexColor('#e2e8f0')), # Horizontal dividers
                        ('ROWBACKGROUNDS', (0, last_row_idx), (-1, last_row_idx), [colors.HexColor('#e0f2fe')]),
                        ('FONTNAME', (0, last_row_idx), (-1, last_row_idx), 'Helvetica-Bold'),
                    ]))
                    story.append(breakdown_table)
                    story.append(Spacer(1, 15))
                
                # Detected Issues
                if strict_eval.get('top_issues'):
                    top_issues = strict_eval.get('top_issues', [])
                    if top_issues and isinstance(top_issues, list):
                        story.append(Paragraph("Issues Detected", heading_style))
                        for issue in top_issues[:10]:  # Limit to 10 issues
                            if issue:
                                story.append(Paragraph(f"• {str(issue)[:200]}", normal_style))
                        story.append(Spacer(1, 12))
                
                # Red Flags Applied
                flags = strict_eval.get('flags', {})
                if isinstance(flags, dict) and any(flags.values()):
                    story.append(Paragraph("Score Caps Applied", heading_style))
                    if flags.get('nonsense_cap_applied'):
                        story.append(Paragraph("• High nonsense word rate → Max score capped at 70", normal_style))
                    if flags.get('low_conf_cap_applied'):
                        story.append(Paragraph("• Low confidence or mumbling detected → Max score capped at 72", normal_style))
                    if flags.get('speed_red_flag'):
                        story.append(Paragraph("• Speaking pace outside acceptable range → Max score capped at 75", normal_style))
                    if flags.get('filler_cap_applied'):
                        story.append(Paragraph("• Excessive filler words → Max score capped at 78", normal_style))
                    story.append(Spacer(1, 12))
                
                # Detailed Metrics (only if metrics exist and are valid)
                metrics = strict_eval.get('metrics', {})
                
                # FALLBACK: If strict metrics are empty/zero but we have general metrics, use those
                speaking_metrics = analysis_data.get('speaking_metrics', {})
                if not metrics.get('words') and speaking_metrics.get('total_words'):
                    metrics['words'] = speaking_metrics.get('total_words', 0)
                    metrics['duration_sec'] = speaking_metrics.get('total_duration', 0)
                    metrics['wpm'] = speaking_metrics.get('speaking_rate_wpm', 0)
                    pause_summary = analysis_data.get('pause_summary', {})
                    if pause_summary:
                        metrics['pause_count'] = pause_summary.get('total_pauses', 0)
                        metrics['mean_pause_s'] = pause_summary.get('avg_pause_duration', 0)
                
                if metrics and isinstance(metrics, dict):
                    has_metrics = any(metrics.get(key, 0) > 0 for key in ['duration_sec', 'words', 'wpm', 'pause_count'])
                    
                    if has_metrics:
                        story.append(Paragraph("Detailed Metrics", heading_style))
                        
                        metrics_data = [['Metric', 'Value']]
                        
                        # Add metrics only if they exist
                        if metrics.get('duration_sec', 0) > 0:
                            metrics_data.append(['Duration', f"{metrics.get('duration_sec', 0):.1f}s"])
                        if metrics.get('words', 0) > 0:
                            metrics_data.append(['Total Words', str(metrics.get('words', 0))])
                        if metrics.get('wpm', 0) > 0:
                            metrics_data.append(['Words Per Minute (WPM)', f"{metrics.get('wpm', 0):.1f}"])
                        if 'articulation_wpm' in metrics:
                            metrics_data.append(['Articulation WPM', f"{metrics.get('articulation_wpm', 0):.1f}"])
                        if 'pause_count' in metrics:
                            metrics_data.append(['Pause Count', str(metrics.get('pause_count', 0))])
                        if 'mean_pause_s' in metrics:
                            metrics_data.append(['Mean Pause Duration', f"{metrics.get('mean_pause_s', 0):.2f}s"])
                        if 'pause_percent' in metrics:
                            pause_percent = metrics.get('pause_percent', 0)
                            if isinstance(pause_percent, (int, float)):
                                metrics_data.append(['Pause Percent of Total', f"{pause_percent * 100:.1f}%"])
                        if 'filler_per_min' in metrics:
                            metrics_data.append(['Filler Words/min', f"{metrics.get('filler_per_min', 0):.2f}"])
                        if 'low_confidence_rate' in metrics:
                            metrics_data.append(['Low Confidence Rate', f"{metrics.get('low_confidence_rate', 0):.1f}%"])
                        if 'mean_confidence' in metrics:
                            metrics_data.append(['Mean Confidence', f"{metrics.get('mean_confidence', 1.0):.3f}"])
                        if 'mumble_bursts' in metrics:
                            metrics_data.append(['Mumble Bursts', str(metrics.get('mumble_bursts', 0))])
                        if 'lexical_diversity_ttr' in metrics:
                            metrics_data.append(['Lexical Diversity (TTR)', f"{metrics.get('lexical_diversity_ttr', 0):.3f}"])
                        if 'nonsense_rate' in metrics:
                            metrics_data.append(['Nonsense Rate', f"{metrics.get('nonsense_rate', 0):.2f}%"])
                        
                        if len(metrics_data) > 1:  # Only create table if we have data
                            metrics_table = Table(metrics_data, colWidths=[2.5*inch, 2*inch])
                            metrics_table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), primary_color),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                                ('TOPPADDING', (0, 0), (-1, 0), 10),
                                ('BACKGROUND', (0, 1), (-1, -1), bg_light),
                                ('LINEBELOW', (0, 0), (-1, -2), 1, colors.HexColor('#e2e8f0')),
                            ]))
                            story.append(metrics_table)
                            story.append(Spacer(1, 20))
            
            # Full Transcription Section - Modern design
            transcript_heading_style = ParagraphStyle(
                'TranscriptHeading',
                parent=heading_style,
                fontSize=14,
                textColor=primary_color,
                spaceBefore=24,
                spaceAfter=12
            )
            story.append(Paragraph("Full Transcription", transcript_heading_style))
            
            # Get both transcript and enhanced transcript
            transcript = analysis_data.get('transcript', '')
            enhanced_transcript = analysis_data.get('enhanced_transcript', '')
            full_transcript = enhanced_transcript if enhanced_transcript else transcript
            
            # Transcript box style - simpler without border properties that might cause issues
            transcript_box_style = ParagraphStyle(
                'TranscriptBox',
                parent=styles['Normal'],
                fontSize=11,
                leftIndent=15,
                rightIndent=15,
                spaceAfter=8,
                backColor=colors.HexColor('#f9fafb'),
                leading=16
            )
            
            # Wrap transcript in a bordered box
            if full_transcript:
                # Truncate very long transcripts to prevent PDF generation issues
                max_transcript_length = 50000  # Limit to prevent recursion errors
                if len(full_transcript) > max_transcript_length:
                    full_transcript = full_transcript[:max_transcript_length] + "\n\n[Transcript truncated for PDF generation]"
                
                # Split into paragraphs for better readability
                transcript_paragraphs = full_transcript.split('\n') if '\n' in full_transcript else [full_transcript]
                
                # Limit number of paragraphs to prevent PDF issues
                max_paragraphs = 200
                if len(transcript_paragraphs) > max_paragraphs:
                    transcript_paragraphs = transcript_paragraphs[:max_paragraphs]
                    transcript_paragraphs.append("\n[Additional paragraphs truncated]")
                
                for para in transcript_paragraphs:
                    if para.strip():
                        try:
                            # Escape HTML special characters and wrap in Courier font
                            para_escaped = str(para.strip())
                            para_escaped = para_escaped.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            para_escaped = para_escaped.replace('"', '&quot;').replace("'", '&#39;')
                            # Remove any problematic characters that might break PDF generation
                            import re
                            # Keep only printable ASCII and common Unicode characters, remove control chars
                            para_escaped = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', para_escaped)
                            # Limit paragraph length to prevent issues
                            if len(para_escaped) > 1000:
                                para_escaped = para_escaped[:1000] + "..."
                            story.append(Paragraph(f"<font face='Courier'>{para_escaped}</font>", transcript_box_style))
                        except Exception as para_error:
                            logger.warning(f"Could not add transcript paragraph: {para_error}")
                            # Skip problematic paragraphs
                            continue
            else:
                story.append(Paragraph("No transcription available", transcript_box_style))
            
            story.append(Spacer(1, 20))
            
            # Presentation Errors Section - Clean modern format
            error_heading_style = ParagraphStyle(
                'ErrorHeading',
                parent=heading_style,
                fontSize=16,
                textColor=colors.HexColor('#dc2626'),
                spaceBefore=15,
                spaceAfter=12
            )
            story.append(Paragraph("⚠️ Identified Issues", error_heading_style))
            
            # Collect all errors
            errors_list = []
            
            # Filler words errors
            filler_analysis = analysis_data.get('filler_analysis', {})
            filler_count = filler_analysis.get('total_fillers', 0)
            filler_breakdown = filler_analysis.get('filler_breakdown', {})
            if filler_count > 0:
                top_fillers = sorted(filler_breakdown.items(), key=lambda x: x[1], reverse=True)[:3]
                filler_text = ", ".join([f"'{word}' ({count}x)" for word, count in top_fillers])
                errors_list.append({
                    'type': 'Filler Words',
                    'severity': 'High' if filler_count > 10 else 'Medium' if filler_count > 5 else 'Low',
                    'description': f"Used {filler_count} filler words: {filler_text}",
                    'recommendation': 'Practice eliminating filler words by pausing instead'
                })
            
            # Long pauses errors
            pause_analysis = analysis_data.get('pause_analysis', {})
            long_pauses = pause_analysis.get('long_pauses', 0)
            longest_pause = pause_analysis.get('longest_pause', 0)
            if long_pauses > 0:
                errors_list.append({
                    'type': 'Long Pauses',
                    'severity': 'High' if longest_pause > 3 else 'Medium',
                    'description': f"Detected {long_pauses} long pauses (>2s). Longest pause: {longest_pause:.1f}s",
                    'recommendation': 'Practice smoother transitions to reduce awkward silence'
                })
            
            # Low confidence/mumbling
            strict_eval = analysis_data.get('strict_evaluation', {})
            if strict_eval:
                metrics = strict_eval.get('metrics', {})
                low_conf_rate = metrics.get('low_confidence_rate', 0)
                mumble_bursts = metrics.get('mumble_bursts', 0)
                if low_conf_rate > 10 or mumble_bursts > 0:
                    errors_list.append({
                        'type': 'Clarity Issues',
                        'severity': 'High' if low_conf_rate > 20 else 'Medium',
                        'description': f"Low confidence rate: {low_conf_rate:.1f}%. Mumble bursts: {mumble_bursts}",
                        'recommendation': 'Enunciate clearly and speak with more confidence'
                    })
            
            # Speaking pace issues
            speaking_rate = analysis_data.get('speaking_rate_wpm', 0) or analysis_data.get('speaking_metrics', {}).get('speaking_rate_wpm', 0) or analysis_data.get('speaking_metrics', {}).get('wpm', 0)
            if speaking_rate > 0:
                if speaking_rate < 120:
                    errors_list.append({
                        'type': 'Speaking Pace',
                        'severity': 'Medium',
                        'description': f"Speaking too slowly ({speaking_rate:.0f} WPM). Ideal: 140-160 WPM",
                        'recommendation': 'Practice speaking at a faster, more natural pace'
                    })
                elif speaking_rate > 180:
                    errors_list.append({
                        'type': 'Speaking Pace',
                        'severity': 'Medium',
                        'description': f"Speaking too quickly ({speaking_rate:.0f} WPM). Ideal: 140-160 WPM",
                        'recommendation': 'Slow down to improve clarity and comprehension'
                    })
            
            # Grammar/coherence issues from strict evaluation
            if strict_eval and 'top_issues' in strict_eval:
                for issue in strict_eval['top_issues']:
                    if 'grammar' in issue.lower() or 'coherence' in issue.lower():
                        errors_list.append({
                            'type': 'Grammar/Coherence',
                            'severity': 'Medium',
                            'description': issue,
                            'recommendation': 'Review sentence structure and logical flow'
                        })
            
            # Display errors in a cool table format
            if errors_list:
                # Create styles for error table cells
                error_cell_style = ParagraphStyle(
                    'ErrorCell',
                    parent=styles['Normal'],
                    fontSize=9,
                    leading=11,
                    leftIndent=0,
                    rightIndent=0,
                    spaceBefore=0,
                    spaceAfter=0,
                    wordWrap='LTR',
                    alignment=0  # LEFT
                )
                
                error_header_style = ParagraphStyle(
                    'ErrorHeader',
                    parent=styles['Normal'],
                    fontSize=10,
                    fontName='Helvetica-Bold',
                    textColor=colors.white,
                    alignment=0
                )
                
                error_data = [[
                    Paragraph('Error Type', error_header_style),
                    Paragraph('Severity', error_header_style),
                    Paragraph('Description', error_header_style),
                    Paragraph('Recommendation', error_header_style)
                ]]
                
                for error in errors_list:
                    severity_color = {
                        'High': '#ef4444',
                        'Medium': '#f59e0b',
                        'Low': '#3b82f6'
                    }.get(error['severity'], '#6b7280')
                    
                    # Escape HTML in descriptions and recommendations and create Paragraph objects
                    try:
                        import re
                        error_type = str(error.get('type', 'Unknown'))
                        error_type = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', error_type)
                        
                        # Escape HTML for plain text cells
                        desc = str(error.get('description', ''))
                        desc = desc.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        desc = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', desc)
                        if len(desc) > 200:
                            desc = desc[:197] + '...'
                        
                        rec = str(error.get('recommendation', ''))
                        rec = rec.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        rec = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', rec)
                        if len(rec) > 150:
                            rec = rec[:147] + '...'
                        
                        severity_text = error.get('severity', 'Unknown')
                        
                        # Create Paragraph objects for each cell
                        error_data.append([
                            Paragraph(error_type, error_cell_style),
                            Paragraph(f"<font color='{severity_color}'><b>{severity_text}</b></font>", error_cell_style),
                            Paragraph(desc, error_cell_style),
                            Paragraph(rec, error_cell_style)
                        ])
                    except Exception as error_item_error:
                        logger.warning(f"Could not process error item: {error_item_error}")
                        continue
                
                # Improved column widths for better spacing
                error_table = Table(error_data, colWidths=[1.4*inch, 1.0*inch, 2.8*inch, 2.0*inch])
                error_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#ef4444')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('TOPPADDING', (0, 1), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#fef2f2')),
                    ('LINEBELOW', (0, 0), (-1, -2), 1, colors.HexColor('#fca5a5')),
                    ('ROWBACKGROUNDS', (0, 2), (-1, -1), [colors.HexColor('#ffffff'), colors.HexColor('#fef2f2')]),
                ]))
                story.append(error_table)
            else:
                story.append(Paragraph("<b>No significant errors detected!</b>", ParagraphStyle('Success', parent=normal_style, textColor=success_color)))
            
            story.append(Spacer(1, 25))
            
            # Strengths and Improvements
            strengths = analysis_data.get('strengths', [])
            if strengths:
                story.append(Paragraph("Strengths", heading_style))
                for strength in strengths:
                    # Sanitize strength text - remove HTML chars and leading bullets/symbols
                    strength_text = str(strength).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    import re
                    # Remove leading non-alphanumeric characters (except parens/quotes if needed, but simple is safer)
                    strength_text = re.sub(r'^[\s\W_]+', '', strength_text).strip()
                    story.append(Paragraph(f"• {strength_text}", normal_style))
                story.append(Spacer(1, 12))
            
            improvements = analysis_data.get('improvements', [])
            if improvements:
                story.append(Paragraph("Areas for Improvement", heading_style))
                for improvement in improvements:
                    # Sanitize improvement text
                    improvement_text = str(improvement).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    import re
                    # Remove leading non-alphanumeric characters (except parens/quotes if needed, but simple is safer)
                    improvement_text = re.sub(r'^[\s\W_]+', '', improvement_text).strip()
                    story.append(Paragraph(f"• {improvement_text}", normal_style))
                story.append(Spacer(1, 20))
            
            # Generate charts
            chart_path = self._generate_score_chart(session_id, analysis_data)
            if chart_path and chart_path.exists():
                story.append(Paragraph("Performance Analytics", heading_style))
                try:
                    story.append(Image(str(chart_path), width=5*inch, height=3*inch))
                except Exception as img_error:
                    logger.warning(f"Could not add chart image: {img_error}")
                    story.append(Paragraph("Chart unavailable", normal_style))
            
            # Build PDF with comprehensive error handling
            try:
                # Validate story has content
                if not story:
                    logger.error("PDF story is empty, cannot generate PDF")
                    raise ValueError("PDF story is empty")
                
                # Ensure directory exists
                pdf_path.parent.mkdir(parents=True, exist_ok=True)
                
                doc.build(story)
                
                # Verify file was created
                if not pdf_path.exists():
                    logger.error(f"PDF file was not created at {pdf_path}")
                    raise FileNotFoundError(f"PDF file was not created at {pdf_path}")
                
                # Check file size (should be > 0)
                if pdf_path.stat().st_size == 0:
                    logger.error(f"PDF file is empty at {pdf_path}")
                    raise ValueError(f"PDF file is empty at {pdf_path}")
                
                logger.info(f"PDF report generated successfully: {pdf_path} (size: {pdf_path.stat().st_size} bytes)")
                return pdf_path
            except Exception as build_error:
                logger.error(f"Error building PDF document: {build_error}", exc_info=True)
                # Try to build a simpler fallback version if full build fails
                try:
                    logger.info("Attempting to generate simplified PDF report...")
                    simple_pdf_path = self.reports_dir / f"{session_id}_simple_report.pdf"
                    simple_pdf_path.parent.mkdir(parents=True, exist_ok=True)
                    simple_doc = SimpleDocTemplate(str(simple_pdf_path), pagesize=letter)
                    
                    # Get transcript for fallback
                    transcript_fallback = analysis_data.get('transcript', '') or analysis_data.get('enhanced_transcript', '')
                    
                    simple_story = [
                        Paragraph("Face2Phase Communication Analysis", title_style),
                        Spacer(1, 12),
                        Paragraph(f"Session ID: {session_id}", styles['Normal']),
                        Spacer(1, 20),
                        Paragraph(f"Overall Score: {overall_score:.1f}/100", heading_style),
                        Spacer(1, 15),
                    ]
                    
                    # Add basic scores
                    simple_scores = [
                        ['Metric', 'Score'],
                        ['Voice Confidence', f"{analysis_data.get('voice_confidence', 0):.1f}/100"],
                        ['Vocabulary Score', f"{analysis_data.get('vocabulary_score', 0):.1f}/100"],
                    ]
                    if not analysis_data.get('is_audio_only', False):
                        simple_scores.insert(2, ['Facial Confidence', f"{analysis_data.get('facial_confidence', 0):.1f}/100"])
                    
                    simple_table = Table(simple_scores, colWidths=[3*inch, 2*inch])
                    simple_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8b5cf6')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ]))
                    simple_story.append(simple_table)
                    simple_story.append(Spacer(1, 20))
                    
                    # Add transcript if available (truncated)
                    if transcript_fallback:
                        try:
                            simple_story.append(Paragraph("Full Transcription", heading_style))
                            transcript_preview = str(transcript_fallback)[:2000]
                            transcript_preview = transcript_preview.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            transcript_preview = transcript_preview.replace('"', '&quot;').replace("'", '&#39;')
                            # Remove control characters
                            import re
                            transcript_preview = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', transcript_preview)
                            if len(str(transcript_fallback)) > 2000:
                                transcript_preview += "..."
                            simple_story.append(Paragraph(f"<font face='Courier'>{transcript_preview}</font>", styles['Normal']))
                        except Exception as transcript_error:
                            logger.warning(f"Could not add transcript to simplified PDF: {transcript_error}")
                            simple_story.append(Paragraph("Transcript available in full report", styles['Normal']))
                    
                    simple_doc.build(simple_story)
                    
                    # Verify simplified PDF was created
                    if not simple_pdf_path.exists() or simple_pdf_path.stat().st_size == 0:
                        logger.error(f"Simplified PDF file was not created or is empty at {simple_pdf_path}")
                        return None
                    
                    logger.info(f"Simplified PDF report generated: {simple_pdf_path} (size: {simple_pdf_path.stat().st_size} bytes)")
                    return simple_pdf_path
                except Exception as simple_error:
                    logger.error(f"Error generating simplified PDF: {simple_error}", exc_info=True)
                    return None
            
        except Exception as e:
            logger.error(f"Error generating PDF report: {e}", exc_info=True)
            return None
    
    def generate_word_report(self, session_id: str, analysis_data: Dict) -> Path:
        """
        Generate comprehensive Word document report
        
        Args:
            session_id: Session identifier
            analysis_data: Complete analysis results
            
        Returns:
            Path to generated Word file
        """
        try:
            word_path = self.reports_dir / f"{session_id}_report.docx"
            
            # Create Word document
            doc = Document()
            
            # Title
            title = doc.add_heading('Face2Phrase Pro - Communication Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Session Info
            doc.add_paragraph(f"Session ID: {session_id}")
            doc.add_paragraph(f"Date: {analysis_data.get('timestamp', datetime.now().isoformat())}")
            doc.add_paragraph()
            
            # Overall Score
            doc.add_heading('Overall Communication Score', level=1)
            score_para = doc.add_paragraph(f"{analysis_data.get('overall_score', 0):.1f}/100")
            score_para.runs[0].font.size = Pt(36)
            score_para.runs[0].font.color.rgb = RGBColor(46, 134, 171)
            score_para.runs[0].bold = True
            doc.add_paragraph()
            
            # Detailed Scores
            doc.add_heading('Detailed Scores', level=1)
            table = doc.add_table(rows=4, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Score'
            hdr_cells[2].text = 'Status'
            
            # Data rows (conditionally hide facial confidence for audio-only)
            metrics = [
                ('Voice Confidence', analysis_data.get('voice_confidence', 0)),
            ]
            
            # Only add facial confidence if NOT audio-only
            if not analysis_data.get('is_audio_only', False):
                metrics.append(('Facial Confidence', analysis_data.get('facial_confidence', 0)))
            
            metrics.append(('Vocabulary Score', analysis_data.get('vocabulary_score', 0)))
            
            for i, (metric, score) in enumerate(metrics, 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = metric
                row_cells[1].text = f"{score:.1f}/100"
                row_cells[2].text = self._get_status(score)
            
            doc.add_paragraph()
            
            # Enhanced Transcript
            doc.add_heading('Enhanced Transcript with Analysis', level=1)
            enhanced_transcript = analysis_data.get('enhanced_transcript', '')
            
            transcript_para = doc.add_paragraph()
            transcript_para.style = 'Intense Quote'
            
            for line in enhanced_transcript.split('\n'):
                doc.add_paragraph(line, style='Code')
            
            doc.add_paragraph()
            
            # Pause Analysis
            pause_analysis = analysis_data.get('pause_analysis', {})
            if pause_analysis:
                doc.add_heading('Pause Analysis', level=1)
                
                pause_table = doc.add_table(rows=6, cols=2)
                pause_table.style = 'Medium Shading 1 Accent 2'
                
                pause_metrics = [
                    ('Total Pauses', str(pause_analysis.get('total_pauses', 0))),
                    ('Long Pauses (>2s)', str(pause_analysis.get('long_pauses', 0))),
                    ('Short Pauses (0.5-2s)', str(pause_analysis.get('short_pauses', 0))),
                    ('Total Silence Time', f"{pause_analysis.get('total_silence_time', 0):.1f}s"),
                    ('Silence Percentage', f"{pause_analysis.get('silence_percentage', 0):.1f}%"),
                ]
                
                for i, (metric, value) in enumerate(pause_metrics):
                    row_cells = pause_table.rows[i].cells
                    row_cells[0].text = metric
                    row_cells[1].text = value
                
                doc.add_paragraph()
            
            # Strengths
            strengths = analysis_data.get('strengths', [])
            if strengths:
                doc.add_heading('Strengths', level=1)
                for strength in strengths:
                    doc.add_paragraph(strength, style='List Bullet')
                doc.add_paragraph()
            
            # Improvements
            improvements = analysis_data.get('improvements', [])
            if improvements:
                doc.add_heading('Areas for Improvement', level=1)
                for improvement in improvements:
                    doc.add_paragraph(improvement, style='List Bullet')
                doc.add_paragraph()
            
            # Add chart
            chart_path = self._generate_score_chart(session_id, analysis_data)
            if chart_path and chart_path.exists():
                doc.add_heading('Performance Analytics', level=1)
                doc.add_picture(str(chart_path), width=Inches(6))
            
            # Save document
            doc.save(str(word_path))
            logger.info(f"Word report generated: {word_path}")
            
            return word_path
            
        except Exception as e:
            logger.error(f"Error generating Word report: {e}")
            return None
    
    def _generate_score_chart(self, session_id: str, analysis_data: Dict) -> Path:
        """Generate modern score visualization chart"""
        try:
            chart_path = self.charts_dir / f"{session_id}_scores.png"
            
            # Set modern style
            plt.style.use('seaborn-v0_8-whitegrid')
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
            fig.patch.set_facecolor('#ffffff')
            
            # Subplot 1: Score Breakdown with gradient-style bars
            categories = ['Voice', 'Facial', 'Vocab']
            scores = [
                analysis_data.get('voice_confidence', 0),
                analysis_data.get('facial_confidence', 0),
                analysis_data.get('vocabulary_score', 0)
            ]
            
            # Modern color palette
            bar_colors = ['#3b82f6', '#8b5cf6', '#06b6d4']
            
            bars = ax1.bar(categories, scores, color=bar_colors, width=0.6, edgecolor='white', linewidth=2)
            ax1.set_ylim(0, 105)
            ax1.set_ylabel('Score', fontsize=11, fontweight='600', color='#374151')
            ax1.set_title('Communication Scores', fontsize=14, fontweight='bold', color='#1e3a5f', pad=15)
            ax1.axhline(y=70, color='#10b981', linestyle='--', alpha=0.6, linewidth=1.5, label='Good (70+)')
            ax1.axhline(y=50, color='#f59e0b', linestyle='--', alpha=0.6, linewidth=1.5, label='Average (50+)')
            ax1.legend(loc='upper right', fontsize=9)
            ax1.set_facecolor('#fafafa')
            ax1.tick_params(colors='#6b7280')
            ax1.spines['top'].set_visible(False)
            ax1.spines['right'].set_visible(False)
            
            # Add value labels on bars with modern styling
            for bar in bars:
                height = bar.get_height()
                ax1.text(bar.get_x() + bar.get_width()/2., height + 2,
                        f'{height:.0f}',
                        ha='center', va='bottom', fontweight='bold', fontsize=12, color='#1e3a5f')
            
            # Subplot 2: Pause Analysis with modern design
            pause_analysis = analysis_data.get('pause_analysis', {})
            pause_categories = ['Long\nPauses', 'Short\nPauses', 'Silence %']
            pause_values = [
                pause_analysis.get('long_pauses', 0),
                pause_analysis.get('short_pauses', 0),
                pause_analysis.get('silence_percentage', 0)
            ]
            
            # Softer color palette for pause chart
            pause_colors = ['#ef4444', '#f59e0b', '#22c55e']
            
            bars2 = ax2.bar(pause_categories, pause_values, color=pause_colors, width=0.6, edgecolor='white', linewidth=2)
            ax2.set_ylabel('Count / %', fontsize=11, fontweight='600', color='#374151')
            ax2.set_title('Pause Analysis', fontsize=14, fontweight='bold', color='#1e3a5f', pad=15)
            ax2.set_facecolor('#fafafa')
            ax2.tick_params(colors='#6b7280')
            ax2.spines['top'].set_visible(False)
            ax2.spines['right'].set_visible(False)
            
            # Add value labels with modern styling
            for bar in bars2:
                height = bar.get_height()
                ax2.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                        f'{height:.1f}',
                        ha='center', va='bottom', fontweight='bold', fontsize=11, color='#1e3a5f')
            
            plt.tight_layout(pad=2.0)
            plt.savefig(str(chart_path), dpi=150, bbox_inches='tight', facecolor='white')
            plt.close()
            
            logger.info(f"Chart generated: {chart_path}")
            return chart_path
            
        except Exception as e:
            logger.error(f"Error generating chart: {e}")
            return None
    
    def _get_status(self, score: float) -> str:
        """Get status label based on score"""
        if score >= 80:
            return "Excellent"
        elif score >= 70:
            return "Good"
        elif score >= 60:
            return "Average"
        elif score >= 50:
            return "Fair"
        else:
            return "Needs Work"

# Global PDF/Word report generator instance
pdf_word_generator = PDFWordReportGenerator()

