"""
PDF and Word Report Generator with Charts and Analytics
"""
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
import json

# PDF Generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

# Word Document Generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Charts
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import logging as mpl_logging
mpl_logging.getLogger('matplotlib.font_manager').setLevel(mpl_logging.WARNING)  # Suppress font manager debug logs
import seaborn as sns
import numpy as np

from config import REPORTS_DIR

logger = logging.getLogger(__name__)

class PDFWordReportGenerator:
    """Generate professional PDF and Word reports with analytics"""
    
    def __init__(self):
        self.reports_dir = REPORTS_DIR
        self.charts_dir = self.reports_dir / "charts"
        self.charts_dir.mkdir(exist_ok=True)
    
    def generate_pdf_report(self, session_id: str, analysis_data: Dict) -> Optional[Path]:
        """
        Generate comprehensive PDF report with charts
        
        Args:
            session_id: Session identifier
            analysis_data: Complete analysis results
            
        Returns:
            Path to generated PDF file or None if generation fails
        """
        try:
            # Ensure reports directory exists
            self.reports_dir.mkdir(exist_ok=True)
            pdf_path = self.reports_dir / f"{session_id}_report.pdf"
            
            # Create PDF document
            doc = SimpleDocTemplate(
                str(pdf_path),
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Container for PDF elements
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#2E86AB'),
                spaceAfter=30,
                alignment=1  # Center
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=colors.HexColor('#A23B72'),
                spaceAfter=12,
                spaceBefore=12
            )
            
            story.append(Paragraph("Face2Phase Communication Analysis", title_style))
            story.append(Spacer(1, 12))
            
            # Session Info
            timestamp = analysis_data.get('timestamp', datetime.now().isoformat())
            info_style = ParagraphStyle(
                'InfoStyle',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.HexColor('#6b7280'),
                alignment=1
            )
            story.append(Paragraph(f"Session ID: {session_id} | Date: {timestamp}", info_style))
            story.append(Spacer(1, 20))
            
            # Overall Score Section with cool design
            overall_score = analysis_data.get('overall_score', 0)
            score_box_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8b5cf6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ROWBACKGROUNDS', (0, 1), (-1, 1), [colors.HexColor('#f3f4f6')]),
            ])
            score_box = Table([['Overall Communication Score'], [f"{overall_score:.1f}/100"]], 
                            colWidths=[5*inch], rowHeights=[0.5*inch, 1*inch])
            score_box.setStyle(score_box_style)
            story.append(score_box)
            story.append(Spacer(1, 20))
            
            # Scores Table (conditionally hide facial confidence for audio-only)
            scores_data = [
                ['Metric', 'Score', 'Status'],
                ['Voice Confidence', f"{analysis_data.get('voice_confidence', 0):.1f}/100", self._get_status(analysis_data.get('voice_confidence', 0))],
            ]
            
            # Only add facial confidence if NOT audio-only
            if not analysis_data.get('is_audio_only', False):
                scores_data.append(['Facial Confidence', f"{analysis_data.get('facial_confidence', 0):.1f}/100", self._get_status(analysis_data.get('facial_confidence', 0))])
            
            scores_data.append(['Vocabulary Score', f"{analysis_data.get('vocabulary_score', 0):.1f}/100", self._get_status(analysis_data.get('vocabulary_score', 0))])
            
            scores_table = Table(scores_data, colWidths=[2.5*inch, 1.5*inch, 1.5*inch])
            scores_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86AB')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(scores_table)
            story.append(Spacer(1, 20))
            
            # Strict Evaluation Breakdown (if available)
            if 'strict_evaluation' in analysis_data:
                strict_eval = analysis_data.get('strict_evaluation', {})
                scores = strict_eval.get('scores', {})
                
                # Only show breakdown if we have the required scores
                if scores and any(key in scores for key in ['clarity_pronunciation_25', 'fluency_pace_20', 'coherence_grammar_25']):
                    story.append(Paragraph("Detailed Score Breakdown", heading_style))
                    
                    breakdown_data = [['Dimension', 'Points', 'Score', 'Max']]
                    
                    # Add each score only if it exists
                    if 'clarity_pronunciation_25' in scores:
                        breakdown_data.append(['Clarity & Pronunciation', '25', f"{scores.get('clarity_pronunciation_25', 0):.1f}", '25.0'])
                    if 'fluency_pace_20' in scores:
                        breakdown_data.append(['Fluency & Pace', '20', f"{scores.get('fluency_pace_20', 0):.1f}", '20.0'])
                    if 'coherence_grammar_25' in scores:
                        breakdown_data.append(['Coherence & Grammar', '25', f"{scores.get('coherence_grammar_25', 0):.1f}", '25.0'])
                    if 'content_accuracy_20' in scores:
                        breakdown_data.append(['Content Accuracy', '20', f"{scores.get('content_accuracy_20', 0):.1f}", '20.0'])
                    if 'delivery_engagement_10' in scores:
                        breakdown_data.append(['Delivery & Engagement', '10', f"{scores.get('delivery_engagement_10', 0):.1f}", '10.0'])
                    if 'final_100' in scores:
                        breakdown_data.append(['Total Score', '100', f"<b>{scores.get('final_100', 0):.1f}</b>", '100.0'])
                
                    if len(breakdown_data) > 1:  # Only create table if we have data rows
                        breakdown_table = Table(breakdown_data, colWidths=[2.2*inch, 1.2*inch, 1.2*inch, 1.2*inch])
                        
                        # Find the last row index for total score highlighting
                        last_row_idx = len(breakdown_data) - 1
                        
                        breakdown_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8b5cf6')),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 11),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f9fafb')),
                            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
                            ('ROWBACKGROUNDS', (0, last_row_idx), (-1, last_row_idx), [colors.HexColor('#F8C471')]),
                            ('FONTNAME', (0, last_row_idx), (-1, last_row_idx), 'Helvetica-Bold'),
                        ]))
                        story.append(breakdown_table)
                        story.append(Spacer(1, 15))
                
                # Detected Issues
                if strict_eval.get('top_issues'):
                    top_issues = strict_eval.get('top_issues', [])
                    if top_issues and isinstance(top_issues, list):
                        story.append(Paragraph("Issues Detected", heading_style))
                        for issue in top_issues[:10]:  # Limit to 10 issues
                            if issue:
                                story.append(Paragraph(f"• {str(issue)[:200]}", styles['Normal']))
                        story.append(Spacer(1, 12))
                
                # Red Flags Applied
                flags = strict_eval.get('flags', {})
                if any(flags.values()):
                    story.append(Paragraph("Score Caps Applied", heading_style))
                    if flags.get('nonsense_cap_applied'):
                        story.append(Paragraph("• High nonsense word rate → Max score capped at 70", styles['Normal']))
                    if flags.get('low_conf_cap_applied'):
                        story.append(Paragraph("• Low confidence or mumbling detected → Max score capped at 72", styles['Normal']))
                    if flags.get('speed_red_flag'):
                        story.append(Paragraph("• Speaking pace outside acceptable range → Max score capped at 75", styles['Normal']))
                    if flags.get('filler_cap_applied'):
                        story.append(Paragraph("• Excessive filler words → Max score capped at 78", styles['Normal']))
                    story.append(Spacer(1, 12))
                
                # Detailed Metrics (only if metrics exist and are valid)
                metrics = strict_eval.get('metrics', {})
                if metrics and isinstance(metrics, dict):
                    # Only show metrics table if we have meaningful data
                    has_metrics = any(key in metrics for key in ['duration_sec', 'words', 'wpm', 'pause_count'])
                    if has_metrics:
                        story.append(Paragraph("Detailed Metrics", heading_style))
                        
                        metrics_data = [['Metric', 'Value']]
                        
                        # Add metrics only if they exist
                        if 'duration_sec' in metrics:
                            metrics_data.append(['Duration', f"{metrics.get('duration_sec', 0):.1f}s"])
                        if 'words' in metrics:
                            metrics_data.append(['Total Words', str(metrics.get('words', 0))])
                        if 'wpm' in metrics:
                            metrics_data.append(['Words Per Minute (WPM)', f"{metrics.get('wpm', 0):.1f}"])
                        if 'articulation_wpm' in metrics:
                            metrics_data.append(['Articulation WPM', f"{metrics.get('articulation_wpm', 0):.1f}"])
                        if 'pause_count' in metrics:
                            metrics_data.append(['Pause Count', str(metrics.get('pause_count', 0))])
                        if 'mean_pause_s' in metrics:
                            metrics_data.append(['Mean Pause Duration', f"{metrics.get('mean_pause_s', 0):.2f}s"])
                        if 'pause_percent' in metrics:
                            pause_percent = metrics.get('pause_percent', 0)
                            if isinstance(pause_percent, (int, float)):
                                metrics_data.append(['Pause Percent of Total', f"{pause_percent * 100:.1f}%"])
                        if 'filler_per_min' in metrics:
                            metrics_data.append(['Filler Words/min', f"{metrics.get('filler_per_min', 0):.2f}"])
                        if 'low_confidence_rate' in metrics:
                            metrics_data.append(['Low Confidence Rate', f"{metrics.get('low_confidence_rate', 0):.1f}%"])
                        if 'mean_confidence' in metrics:
                            metrics_data.append(['Mean Confidence', f"{metrics.get('mean_confidence', 1.0):.3f}"])
                        if 'mumble_bursts' in metrics:
                            metrics_data.append(['Mumble Bursts', str(metrics.get('mumble_bursts', 0))])
                        if 'lexical_diversity_ttr' in metrics:
                            metrics_data.append(['Lexical Diversity (TTR)', f"{metrics.get('lexical_diversity_ttr', 0):.3f}"])
                        if 'nonsense_rate' in metrics:
                            metrics_data.append(['Nonsense Rate', f"{metrics.get('nonsense_rate', 0):.2f}%"])
                        
                        if len(metrics_data) > 1:  # Only create table if we have data
                            metrics_table = Table(metrics_data, colWidths=[2.5*inch, 2*inch])
                            metrics_table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8b5cf6')),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f9fafb')),
                                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb'))
                            ]))
                            story.append(metrics_table)
                            story.append(Spacer(1, 20))
            
            # Full Transcription Section - Prominent
            transcript_heading_style = ParagraphStyle(
                'TranscriptHeading',
                parent=heading_style,
                fontSize=18,
                textColor=colors.HexColor('#8b5cf6'),
                backColor=colors.HexColor('#f3f4f6'),
                borderPadding=10,
                spaceAfter=15
            )
            story.append(Paragraph("Full Transcription", transcript_heading_style))
            
            # Get both transcript and enhanced transcript
            transcript = analysis_data.get('transcript', '')
            enhanced_transcript = analysis_data.get('enhanced_transcript', '')
            full_transcript = enhanced_transcript if enhanced_transcript else transcript
            
            # Transcript box style - simpler without border properties that might cause issues
            transcript_box_style = ParagraphStyle(
                'TranscriptBox',
                parent=styles['Normal'],
                fontSize=11,
                leftIndent=15,
                rightIndent=15,
                spaceAfter=8,
                backColor=colors.HexColor('#f9fafb'),
                leading=16
            )
            
            # Wrap transcript in a bordered box
            if full_transcript:
                # Truncate very long transcripts to prevent PDF generation issues
                max_transcript_length = 50000  # Limit to prevent recursion errors
                if len(full_transcript) > max_transcript_length:
                    full_transcript = full_transcript[:max_transcript_length] + "\n\n[Transcript truncated for PDF generation]"
                
                # Split into paragraphs for better readability
                transcript_paragraphs = full_transcript.split('\n') if '\n' in full_transcript else [full_transcript]
                
                # Limit number of paragraphs to prevent PDF issues
                max_paragraphs = 200
                if len(transcript_paragraphs) > max_paragraphs:
                    transcript_paragraphs = transcript_paragraphs[:max_paragraphs]
                    transcript_paragraphs.append("\n[Additional paragraphs truncated]")
                
                for para in transcript_paragraphs:
                    if para.strip():
                        try:
                            # Escape HTML special characters and wrap in Courier font
                            para_escaped = str(para.strip())
                            para_escaped = para_escaped.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            para_escaped = para_escaped.replace('"', '&quot;').replace("'", '&#39;')
                            # Remove any problematic characters that might break PDF generation
                            import re
                            # Keep only printable ASCII and common Unicode characters, remove control chars
                            para_escaped = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', para_escaped)
                            # Limit paragraph length to prevent issues
                            if len(para_escaped) > 1000:
                                para_escaped = para_escaped[:1000] + "..."
                            story.append(Paragraph(f"<font face='Courier'>{para_escaped}</font>", transcript_box_style))
                        except Exception as para_error:
                            logger.warning(f"Could not add transcript paragraph: {para_error}")
                            # Skip problematic paragraphs
                            continue
            else:
                story.append(Paragraph("No transcription available", transcript_box_style))
            
            story.append(Spacer(1, 25))
            
            # Presentation Errors Section - Cool Format
            error_heading_style = ParagraphStyle(
                'ErrorHeading',
                parent=heading_style,
                fontSize=18,
                textColor=colors.HexColor('#ef4444'),
                spaceAfter=15
            )
            story.append(Paragraph("Presentation Errors & Issues", error_heading_style))
            
            # Collect all errors
            errors_list = []
            
            # Filler words errors
            filler_analysis = analysis_data.get('filler_analysis', {})
            filler_count = filler_analysis.get('total_fillers', 0)
            filler_breakdown = filler_analysis.get('filler_breakdown', {})
            if filler_count > 0:
                top_fillers = sorted(filler_breakdown.items(), key=lambda x: x[1], reverse=True)[:3]
                filler_text = ", ".join([f"'{word}' ({count}x)" for word, count in top_fillers])
                errors_list.append({
                    'type': 'Filler Words',
                    'severity': 'High' if filler_count > 10 else 'Medium' if filler_count > 5 else 'Low',
                    'description': f"Used {filler_count} filler words: {filler_text}",
                    'recommendation': 'Practice eliminating filler words by pausing instead'
                })
            
            # Long pauses errors
            pause_analysis = analysis_data.get('pause_analysis', {})
            long_pauses = pause_analysis.get('long_pauses', 0)
            longest_pause = pause_analysis.get('longest_pause', 0)
            if long_pauses > 0:
                errors_list.append({
                    'type': 'Long Pauses',
                    'severity': 'High' if longest_pause > 3 else 'Medium',
                    'description': f"Detected {long_pauses} long pauses (>2s). Longest pause: {longest_pause:.1f}s",
                    'recommendation': 'Practice smoother transitions to reduce awkward silence'
                })
            
            # Low confidence/mumbling
            strict_eval = analysis_data.get('strict_evaluation', {})
            if strict_eval:
                metrics = strict_eval.get('metrics', {})
                low_conf_rate = metrics.get('low_confidence_rate', 0)
                mumble_bursts = metrics.get('mumble_bursts', 0)
                if low_conf_rate > 10 or mumble_bursts > 0:
                    errors_list.append({
                        'type': 'Clarity Issues',
                        'severity': 'High' if low_conf_rate > 20 else 'Medium',
                        'description': f"Low confidence rate: {low_conf_rate:.1f}%. Mumble bursts: {mumble_bursts}",
                        'recommendation': 'Enunciate clearly and speak with more confidence'
                    })
            
            # Speaking pace issues
            speaking_rate = analysis_data.get('speaking_rate_wpm', 0) or analysis_data.get('speaking_metrics', {}).get('wpm', 0)
            if speaking_rate > 0:
                if speaking_rate < 120:
                    errors_list.append({
                        'type': 'Speaking Pace',
                        'severity': 'Medium',
                        'description': f"Speaking too slowly ({speaking_rate:.0f} WPM). Ideal: 140-160 WPM",
                        'recommendation': 'Practice speaking at a faster, more natural pace'
                    })
                elif speaking_rate > 180:
                    errors_list.append({
                        'type': 'Speaking Pace',
                        'severity': 'Medium',
                        'description': f"Speaking too quickly ({speaking_rate:.0f} WPM). Ideal: 140-160 WPM",
                        'recommendation': 'Slow down to improve clarity and comprehension'
                    })
            
            # Grammar/coherence issues from strict evaluation
            if strict_eval and 'top_issues' in strict_eval:
                for issue in strict_eval['top_issues']:
                    if 'grammar' in issue.lower() or 'coherence' in issue.lower():
                        errors_list.append({
                            'type': 'Grammar/Coherence',
                            'severity': 'Medium',
                            'description': issue,
                            'recommendation': 'Review sentence structure and logical flow'
                        })
            
            # Display errors in a cool table format
            if errors_list:
                error_data = [['Error Type', 'Severity', 'Description', 'Recommendation']]
                for error in errors_list:
                    severity_color = {
                        'High': '#ef4444',
                        'Medium': '#f59e0b',
                        'Low': '#3b82f6'
                    }.get(error['severity'], '#6b7280')
                    
                    # Escape HTML in descriptions and recommendations
                    try:
                        error_type = str(error.get('type', 'Unknown'))
                        desc = str(error.get('description', '')).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        rec = str(error.get('recommendation', '')).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        # Remove control characters
                        import re
                        error_type = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', error_type)
                        desc = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', desc)
                        rec = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', rec)
                        
                        error_data.append([
                            error_type,
                            f"<font color='{severity_color}'><b>{error.get('severity', 'Unknown')}</b></font>",
                            desc[:200] if len(desc) > 200 else desc,  # Truncate long descriptions
                            rec[:150] if len(rec) > 150 else rec  # Truncate long recommendations
                        ])
                    except Exception as error_item_error:
                        logger.warning(f"Could not process error item: {error_item_error}")
                        continue
                
                error_table = Table(error_data, colWidths=[1.2*inch, 0.8*inch, 2*inch, 1.5*inch])
                error_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#ef4444')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('TOPPADDING', (0, 1), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#fef2f2')),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#fecaca')),
                    ('ROWBACKGROUNDS', (0, 2), (-1, -1), [colors.HexColor('#ffffff'), colors.HexColor('#fef2f2')]),
                ]))
                story.append(error_table)
            else:
                story.append(Paragraph("<font color='#10b981'><b>No significant errors detected!</b></font>", styles['Normal']))
            
            story.append(Spacer(1, 25))
            
            # Strengths and Improvements
            strengths = analysis_data.get('strengths', [])
            if strengths:
                story.append(Paragraph("Strengths", heading_style))
                for strength in strengths:
                    # Sanitize strength text
                    strength_text = str(strength).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(f"• {strength_text}", styles['Normal']))
                story.append(Spacer(1, 12))
            
            improvements = analysis_data.get('improvements', [])
            if improvements:
                story.append(Paragraph("Areas for Improvement", heading_style))
                for improvement in improvements:
                    # Sanitize improvement text
                    improvement_text = str(improvement).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(f"• {improvement_text}", styles['Normal']))
                story.append(Spacer(1, 20))
            
            # Generate charts
            chart_path = self._generate_score_chart(session_id, analysis_data)
            if chart_path and chart_path.exists():
                story.append(Paragraph("Performance Analytics", heading_style))
                try:
                    story.append(Image(str(chart_path), width=5*inch, height=3*inch))
                except Exception as img_error:
                    logger.warning(f"Could not add chart image: {img_error}")
                    story.append(Paragraph("Chart unavailable", styles['Normal']))
            
            # Build PDF with comprehensive error handling
            try:
                # Validate story has content
                if not story:
                    logger.error("PDF story is empty, cannot generate PDF")
                    raise ValueError("PDF story is empty")
                
                # Ensure directory exists
                pdf_path.parent.mkdir(parents=True, exist_ok=True)
                
                doc.build(story)
                
                # Verify file was created
                if not pdf_path.exists():
                    logger.error(f"PDF file was not created at {pdf_path}")
                    raise FileNotFoundError(f"PDF file was not created at {pdf_path}")
                
                # Check file size (should be > 0)
                if pdf_path.stat().st_size == 0:
                    logger.error(f"PDF file is empty at {pdf_path}")
                    raise ValueError(f"PDF file is empty at {pdf_path}")
                
                logger.info(f"PDF report generated successfully: {pdf_path} (size: {pdf_path.stat().st_size} bytes)")
                return pdf_path
            except Exception as build_error:
                logger.error(f"Error building PDF document: {build_error}", exc_info=True)
                # Try to build a simpler fallback version if full build fails
                try:
                    logger.info("Attempting to generate simplified PDF report...")
                    simple_pdf_path = self.reports_dir / f"{session_id}_simple_report.pdf"
                    simple_pdf_path.parent.mkdir(parents=True, exist_ok=True)
                    simple_doc = SimpleDocTemplate(str(simple_pdf_path), pagesize=letter)
                    
                    # Get transcript for fallback
                    transcript_fallback = analysis_data.get('transcript', '') or analysis_data.get('enhanced_transcript', '')
                    
                    simple_story = [
                        Paragraph("Face2Phase Communication Analysis", title_style),
                        Spacer(1, 12),
                        Paragraph(f"Session ID: {session_id}", styles['Normal']),
                        Spacer(1, 20),
                        Paragraph(f"Overall Score: {overall_score:.1f}/100", heading_style),
                        Spacer(1, 15),
                    ]
                    
                    # Add basic scores
                    simple_scores = [
                        ['Metric', 'Score'],
                        ['Voice Confidence', f"{analysis_data.get('voice_confidence', 0):.1f}/100"],
                        ['Vocabulary Score', f"{analysis_data.get('vocabulary_score', 0):.1f}/100"],
                    ]
                    if not analysis_data.get('is_audio_only', False):
                        simple_scores.insert(2, ['Facial Confidence', f"{analysis_data.get('facial_confidence', 0):.1f}/100"])
                    
                    simple_table = Table(simple_scores, colWidths=[3*inch, 2*inch])
                    simple_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8b5cf6')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ]))
                    simple_story.append(simple_table)
                    simple_story.append(Spacer(1, 20))
                    
                    # Add transcript if available (truncated)
                    if transcript_fallback:
                        try:
                            simple_story.append(Paragraph("Full Transcription", heading_style))
                            transcript_preview = str(transcript_fallback)[:2000]
                            transcript_preview = transcript_preview.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            transcript_preview = transcript_preview.replace('"', '&quot;').replace("'", '&#39;')
                            # Remove control characters
                            import re
                            transcript_preview = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', transcript_preview)
                            if len(str(transcript_fallback)) > 2000:
                                transcript_preview += "..."
                            simple_story.append(Paragraph(f"<font face='Courier'>{transcript_preview}</font>", styles['Normal']))
                        except Exception as transcript_error:
                            logger.warning(f"Could not add transcript to simplified PDF: {transcript_error}")
                            simple_story.append(Paragraph("Transcript available in full report", styles['Normal']))
                    
                    simple_doc.build(simple_story)
                    
                    # Verify simplified PDF was created
                    if not simple_pdf_path.exists() or simple_pdf_path.stat().st_size == 0:
                        logger.error(f"Simplified PDF file was not created or is empty at {simple_pdf_path}")
                        return None
                    
                    logger.info(f"Simplified PDF report generated: {simple_pdf_path} (size: {simple_pdf_path.stat().st_size} bytes)")
                    return simple_pdf_path
                except Exception as simple_error:
                    logger.error(f"Error generating simplified PDF: {simple_error}", exc_info=True)
                    return None
            
        except Exception as e:
            logger.error(f"Error generating PDF report: {e}", exc_info=True)
            return None
    
    def generate_word_report(self, session_id: str, analysis_data: Dict) -> Path:
        """
        Generate comprehensive Word document report
        
        Args:
            session_id: Session identifier
            analysis_data: Complete analysis results
            
        Returns:
            Path to generated Word file
        """
        try:
            word_path = self.reports_dir / f"{session_id}_report.docx"
            
            # Create Word document
            doc = Document()
            
            # Title
            title = doc.add_heading('Face2Phrase Pro - Communication Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Session Info
            doc.add_paragraph(f"Session ID: {session_id}")
            doc.add_paragraph(f"Date: {analysis_data.get('timestamp', datetime.now().isoformat())}")
            doc.add_paragraph()
            
            # Overall Score
            doc.add_heading('Overall Communication Score', level=1)
            score_para = doc.add_paragraph(f"{analysis_data.get('overall_score', 0):.1f}/100")
            score_para.runs[0].font.size = Pt(36)
            score_para.runs[0].font.color.rgb = RGBColor(46, 134, 171)
            score_para.runs[0].bold = True
            doc.add_paragraph()
            
            # Detailed Scores
            doc.add_heading('Detailed Scores', level=1)
            table = doc.add_table(rows=4, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Score'
            hdr_cells[2].text = 'Status'
            
            # Data rows (conditionally hide facial confidence for audio-only)
            metrics = [
                ('Voice Confidence', analysis_data.get('voice_confidence', 0)),
            ]
            
            # Only add facial confidence if NOT audio-only
            if not analysis_data.get('is_audio_only', False):
                metrics.append(('Facial Confidence', analysis_data.get('facial_confidence', 0)))
            
            metrics.append(('Vocabulary Score', analysis_data.get('vocabulary_score', 0)))
            
            for i, (metric, score) in enumerate(metrics, 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = metric
                row_cells[1].text = f"{score:.1f}/100"
                row_cells[2].text = self._get_status(score)
            
            doc.add_paragraph()
            
            # Enhanced Transcript
            doc.add_heading('Enhanced Transcript with Analysis', level=1)
            enhanced_transcript = analysis_data.get('enhanced_transcript', '')
            
            transcript_para = doc.add_paragraph()
            transcript_para.style = 'Intense Quote'
            
            for line in enhanced_transcript.split('\n'):
                doc.add_paragraph(line, style='Code')
            
            doc.add_paragraph()
            
            # Pause Analysis
            pause_analysis = analysis_data.get('pause_analysis', {})
            if pause_analysis:
                doc.add_heading('Pause Analysis', level=1)
                
                pause_table = doc.add_table(rows=6, cols=2)
                pause_table.style = 'Medium Shading 1 Accent 2'
                
                pause_metrics = [
                    ('Total Pauses', str(pause_analysis.get('total_pauses', 0))),
                    ('Long Pauses (>2s)', str(pause_analysis.get('long_pauses', 0))),
                    ('Short Pauses (0.5-2s)', str(pause_analysis.get('short_pauses', 0))),
                    ('Total Silence Time', f"{pause_analysis.get('total_silence_time', 0):.1f}s"),
                    ('Silence Percentage', f"{pause_analysis.get('silence_percentage', 0):.1f}%"),
                ]
                
                for i, (metric, value) in enumerate(pause_metrics):
                    row_cells = pause_table.rows[i].cells
                    row_cells[0].text = metric
                    row_cells[1].text = value
                
                doc.add_paragraph()
            
            # Strengths
            strengths = analysis_data.get('strengths', [])
            if strengths:
                doc.add_heading('Strengths', level=1)
                for strength in strengths:
                    doc.add_paragraph(strength, style='List Bullet')
                doc.add_paragraph()
            
            # Improvements
            improvements = analysis_data.get('improvements', [])
            if improvements:
                doc.add_heading('Areas for Improvement', level=1)
                for improvement in improvements:
                    doc.add_paragraph(improvement, style='List Bullet')
                doc.add_paragraph()
            
            # Add chart
            chart_path = self._generate_score_chart(session_id, analysis_data)
            if chart_path and chart_path.exists():
                doc.add_heading('Performance Analytics', level=1)
                doc.add_picture(str(chart_path), width=Inches(6))
            
            # Save document
            doc.save(str(word_path))
            logger.info(f"Word report generated: {word_path}")
            
            return word_path
            
        except Exception as e:
            logger.error(f"Error generating Word report: {e}")
            return None
    
    def _generate_score_chart(self, session_id: str, analysis_data: Dict) -> Path:
        """Generate score visualization chart"""
        try:
            chart_path = self.charts_dir / f"{session_id}_scores.png"
            
            # Create figure with subplots
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
            
            # Subplot 1: Score Breakdown
            categories = ['Voice\nConfidence', 'Facial\nConfidence', 'Vocabulary\nScore']
            scores = [
                analysis_data.get('voice_confidence', 0),
                analysis_data.get('facial_confidence', 0),
                analysis_data.get('vocabulary_score', 0)
            ]
            colors_list = ['#2E86AB', '#A23B72', '#F18F01']
            
            bars = ax1.bar(categories, scores, color=colors_list, alpha=0.7, edgecolor='black')
            ax1.set_ylim(0, 100)
            ax1.set_ylabel('Score', fontsize=12, fontweight='bold')
            ax1.set_title('Communication Scores Breakdown', fontsize=14, fontweight='bold')
            ax1.axhline(y=70, color='green', linestyle='--', alpha=0.5, label='Good (70+)')
            ax1.axhline(y=50, color='orange', linestyle='--', alpha=0.5, label='Average (50+)')
            ax1.legend()
            
            # Add value labels on bars
            for bar in bars:
                height = bar.get_height()
                ax1.text(bar.get_x() + bar.get_width()/2., height,
                        f'{height:.1f}',
                        ha='center', va='bottom', fontweight='bold')
            
            # Subplot 2: Pause Analysis
            pause_analysis = analysis_data.get('pause_analysis', {})
            pause_categories = ['Long\nPauses', 'Short\nPauses', 'Silence\n%']
            pause_values = [
                pause_analysis.get('long_pauses', 0),
                pause_analysis.get('short_pauses', 0),
                pause_analysis.get('silence_percentage', 0)
            ]
            
            bars2 = ax2.bar(pause_categories, pause_values, color=['#C73E1D', '#F59B00', '#6A994E'], alpha=0.7, edgecolor='black')
            ax2.set_ylabel('Count / Percentage', fontsize=12, fontweight='bold')
            ax2.set_title('Pause Analysis', fontsize=14, fontweight='bold')
            
            # Add value labels
            for bar in bars2:
                height = bar.get_height()
                ax2.text(bar.get_x() + bar.get_width()/2., height,
                        f'{height:.1f}',
                        ha='center', va='bottom', fontweight='bold')
            
            plt.tight_layout()
            plt.savefig(str(chart_path), dpi=300, bbox_inches='tight')
            plt.close()
            
            logger.info(f"Chart generated: {chart_path}")
            return chart_path
            
        except Exception as e:
            logger.error(f"Error generating chart: {e}")
            return None
    
    def _get_status(self, score: float) -> str:
        """Get status label based on score"""
        if score >= 80:
            return "Excellent"
        elif score >= 70:
            return "Good"
        elif score >= 60:
            return "Average"
        elif score >= 50:
            return "Fair"
        else:
            return "Needs Work"

# Global PDF/Word report generator instance
pdf_word_generator = PDFWordReportGenerator()

